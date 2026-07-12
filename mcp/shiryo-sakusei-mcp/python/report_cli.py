# -*- coding: utf-8 -*-
"""CLI for 資料作成 MCP — docx operations on security report template."""
import argparse
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
FONT = "ＭＳ ゴシック"
TITLE_HALF_PT = 32


def _set_run_font(run, *, bold=None, size_half_pt=21):
    if bold is not None:
        run.bold = bold
    if size_half_pt is not None:
        run.font.size = Pt(size_half_pt / 2)
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def clear_paragraph(p):
    p._element.clear_content()


def write_simple(p, text, *, bold=False, size_half_pt=21):
    clear_paragraph(p)
    run = p.add_run(text)
    _set_run_font(run, bold=bold, size_half_pt=size_half_pt)


def replace_cell_text(cell, text):
    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    clear_paragraph(p)
    run = p.add_run(text)
    _set_run_font(run)


def _set_txbx_texts(shape_para_xml: str, header: str, body: str) -> str:
    root = ET.fromstring(shape_para_xml)
    for txbx in root.findall(f".//{{{W_NS}}}txbxContent"):
        for child in list(txbx):
            txbx.remove(child)
        p1 = ET.Element(f"{{{W_NS}}}p")
        r1 = ET.SubElement(p1, f"{{{W_NS}}}r")
        rpr1 = ET.SubElement(r1, f"{{{W_NS}}}rPr")
        ET.SubElement(rpr1, f"{{{W_NS}}}b")
        ET.SubElement(rpr1, f"{{{W_NS}}}bCs")
        lang1 = ET.SubElement(rpr1, f"{{{W_NS}}}lang")
        lang1.set(f"{{{W_NS}}}eastAsia", "ja-JP")
        t1 = ET.SubElement(r1, f"{{{W_NS}}}t")
        t1.text = header
        txbx.append(p1)
        p2 = ET.Element(f"{{{W_NS}}}p")
        ppr2 = ET.SubElement(p2, f"{{{W_NS}}}pPr")
        ind = ET.SubElement(ppr2, f"{{{W_NS}}}ind")
        ind.set(f"{{{W_NS}}}firstLineChars", "100")
        ind.set(f"{{{W_NS}}}firstLine", "220")
        r2 = ET.SubElement(p2, f"{{{W_NS}}}r")
        rpr2 = ET.SubElement(r2, f"{{{W_NS}}}rPr")
        lang2 = ET.SubElement(rpr2, f"{{{W_NS}}}lang")
        lang2.set(f"{{{W_NS}}}eastAsia", "ja-JP")
        t2 = ET.SubElement(r2, f"{{{W_NS}}}t")
        t2.text = body
        txbx.append(p2)
    xml = ET.tostring(root, encoding="unicode")
    import random
    uid = f"{random.randint(0, 0xFFFFFFFF):08X}"
    xml = re.sub(r'(anchorId|editId)="[0-9A-F]+"', rf'\1="{uid}"', xml)
    xml = re.sub(r'id="\d+"', f'id="{random.randint(10**9, 10**10 - 1)}"', xml, count=1)
    return xml


def load_shape_xml(assets_dir: Path) -> str:
    shape_path = assets_dir / "roundrect-paragraph.xml"
    if shape_path.is_file():
        return shape_path.read_text(encoding="utf-8")
    raise FileNotFoundError(f"Shape asset missing: {shape_path}")


def insert_definition_box(doc, para_index: int, header: str, body: str, assets_dir: Path):
    shape_xml = _set_txbx_texts(load_shape_xml(assets_dir), header, body)
    new_p = parse_xml(shape_xml)
    target = doc.paragraphs[para_index]
    target._element.getparent().replace(target._element, new_p)


def zen_month(n: int) -> str:
    trans = str.maketrans("0123456789", "０１２３４５６７８９")
    return str(n).translate(trans)


def cmd_set_header(args):
    doc = Document(args.docx)
    title = f"２０２６年{zen_month(args.report_month)}月度情報セキュリティレポート"
    if len(doc.paragraphs) > 0:
        write_simple(doc.paragraphs[0], title, bold=True, size_half_pt=TITLE_HALF_PT)
    if args.meeting_date and len(doc.paragraphs) > 1:
        write_simple(doc.paragraphs[1], args.meeting_date)
    if len(doc.paragraphs) > 2:
        write_simple(doc.paragraphs[2], "経営企画部システム推進室")
    doc.save(args.docx)
    return {"ok": True, "title": title, "path": args.docx}


def cmd_prepare_empty_sections(args):
    doc = Document(args.docx)
    ry, rm = args.report_year, args.report_month
    section_title = f"２.２０２６年{zen_month(rm)}月の情報セキュリティ検知状況"
    period = f"（集計：２０２６年{zen_month(rm)}月１日～{zen_month(rm)}月３０日）"

  # Best-effort paragraph indices from 7月 template structure
    if len(doc.paragraphs) > 38:
        clear_paragraph(doc.paragraphs[38])
        p = doc.paragraphs[38]
        r1 = p.add_run(section_title)
        _set_run_font(r1, bold=True, size_half_pt=TITLE_HALF_PT)
        p.add_run("\n")
        r2 = p.add_run(period)
        _set_run_font(r2)
    if len(doc.paragraphs) > 39:
        write_simple(
            doc.paragraphs[39],
            f"（１）コンピュータウイルス感染について\n　 ・コンピュータウイルス感染被疑事象は",
        )
    if len(doc.paragraphs) > 42:
        clear_paragraph(doc.paragraphs[42])
        p = doc.paragraphs[42]
        for seg in [
            "（２）情報セキュリティ機器の異常検知について\n",
            "・ネットワーク監視（疑わしい通信検知件数）：",
            "　　　",
            "件　",
        ]:
            _set_run_font(p.add_run(seg))
    if len(doc.paragraphs) > 45:
        clear_paragraph(doc.paragraphs[45])
        p = doc.paragraphs[45]
        for seg in ["・ＰＣ端末監視（SKYSEA監視で疑わしい検知）：", "　　　", "件"]:
            _set_run_font(p.add_run(seg))
    if len(doc.paragraphs) > 49:
        write_simple(doc.paragraphs[49], f"２０２６年{rm}月の主なセキュリティ事例（社外一般）　")
    if len(doc.paragraphs) > 50:
        write_simple(doc.paragraphs[50], "\n以上")

    if doc.tables:
        t = doc.tables[-1]
        rows = [
            ("№", "公開日", "分類", "概要", "特徴"),
            ("事例１", "", "", "", ""),
            ("事例２", "", "", "", ""),
            ("事例３", "", "", "", ""),
        ]
        while len(t.rows) < len(rows):
            t.add_row()
        while len(t.rows) > len(rows):
            tbl = t._tbl
            tbl.remove(t.rows[-1]._tr)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                replace_cell_text(t.rows[ri].cells[ci], val)

    doc.save(args.docx)
    return {"ok": True, "path": args.docx, "section": section_title}


def cmd_insert_definition_box(args):
    doc = Document(args.docx)
    assets = Path(args.assets_dir)
    insert_definition_box(doc, args.paragraph_index, args.header, args.body, assets)
    doc.save(args.docx)
    return {"ok": True, "path": args.docx, "header": args.header}


def cmd_extract_text(args):
    doc = Document(args.docx)
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for ti, table in enumerate(doc.tables):
        parts.append(f"[TABLE {ti}]")
        for row in table.rows:
            cells = [c.text.replace("\n", " / ").strip() for c in row.cells]
            parts.append(" | ".join(cells))
    # Text inside shapes/textboxes is not in paragraph.text
    with zipfile.ZipFile(args.docx) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    for m in re.finditer(r"<w:t[^>]*>([^<]+)</w:t>", xml):
        t = m.group(1).strip()
        if t and t not in "\n".join(parts):
            parts.append(t)
    return {"ok": True, "text": "\n".join(parts), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables)}


def main():
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    parser = argparse.ArgumentParser()
    sub = parser.add_subparsers(dest="cmd", required=True)

    p1 = sub.add_parser("set_header")
    p1.add_argument("--docx", required=True)
    p1.add_argument("--report-month", type=int, required=True)
    p1.add_argument("--meeting-date", default="")

    p2 = sub.add_parser("prepare_empty_sections")
    p2.add_argument("--docx", required=True)
    p2.add_argument("--report-year", type=int, default=2026)
    p2.add_argument("--report-month", type=int, required=True)

    p3 = sub.add_parser("insert_definition_box")
    p3.add_argument("--docx", required=True)
    p3.add_argument("--header", required=True)
    p3.add_argument("--body", required=True)
    p3.add_argument("--paragraph-index", type=int, default=8)
    p3.add_argument("--assets-dir", required=True)

    p4 = sub.add_parser("extract_text")
    p4.add_argument("--docx", required=True)

    args = parser.parse_args()
    assets_default = Path(__file__).resolve().parent.parent / "assets"

    if args.cmd == "set_header":
        result = cmd_set_header(args)
    elif args.cmd == "prepare_empty_sections":
        result = cmd_prepare_empty_sections(args)
    elif args.cmd == "insert_definition_box":
        if not hasattr(args, "assets_dir") or not args.assets_dir:
            args.assets_dir = str(assets_default)
        result = cmd_insert_definition_box(args)
    elif args.cmd == "extract_text":
        result = cmd_extract_text(args)
    else:
        result = {"ok": False, "error": "unknown command"}

    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
