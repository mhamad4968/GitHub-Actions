# -*- coding: utf-8 -*-
"""Word DOCX formatting helpers — 4月情報セキュリティレポートテンプレート準拠（R2/R5）."""
from __future__ import annotations

import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

FONT_GOTHIC = "ＭＳ ゴシック"
FONT_MINCHO = "ＭＳ 明朝"


def set_run_font(run, font_name: str, size_pt: float, bold: bool | None = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        if child.tag.endswith("r") or child.tag.endswith("hyperlink"):
            element.remove(child)


def set_paragraph_text(
    paragraph,
    text: str,
    *,
    font_name: str = FONT_GOTHIC,
    size_pt: float = 10.5,
    bold: bool | None = False,
    align=None,
) -> None:
    clear_paragraph(paragraph)
    if not text:
        if align is not None:
            paragraph.alignment = align
        return
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size_pt, bold)
    if align is not None:
        paragraph.alignment = align


def paragraph_format_for(text: str) -> dict:
    """Classify paragraph text. Prefix rules use **raw** text (R5: strip before prefix breaks 　・)."""
    raw = text
    t = text.strip()
    base = {"font_name": FONT_GOTHIC, "size_pt": 10.5, "bold": False, "align": None}
    if not t:
        return {**base, "empty": True}
    if re.fullmatch(r"２０２６年.+情報セキュリティレポート", t):
        return {**base, "size_pt": 16, "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER}
    if t == "経営企画部システム推進室" or re.fullmatch(r"2026年\d+月\d+日", t):
        return {**base, "align": WD_ALIGN_PARAGRAPH.RIGHT}
    if t.startswith("表題の件"):
        return {**base, "size_pt": 12}
    if t.startswith("１.") or (t.startswith("２.") and "情報セキュリティ検知" in t):
        return {**base, "size_pt": 16, "bold": True}
    if t.startswith("【注意喚起】") or t.startswith("～"):
        return {**base, "bold": True}
    if t.startswith("（集計："):
        return base
    if t.startswith("（１）") or t.startswith("（２）"):
        return {**base, "size_pt": 12, "bold": True}
    if "主なセキュリティ事例" in t:
        return {**base, "size_pt": 12, "bold": True}
    if t == "以上":
        return {**base, "align": WD_ALIGN_PARAGRAPH.RIGHT}
    if t.startswith("●"):
        return {**base, "bold": True}
    if raw.startswith("　 ・"):
        return {**base, "bold": True}
    if t.startswith("・ネットワーク監視") or t.startswith("・ＰＣ端末監視"):
        return {**base, "bold": True}
    if t.startswith("・"):
        return base
    if t.startswith("（出典") or t == "（下表参照）":
        return base
    return base


def apply_paragraph_format(paragraph) -> None:
    text = paragraph.text
    spec = paragraph_format_for(text)
    if spec.get("empty"):
        clear_paragraph(paragraph)
        return
    set_paragraph_text(paragraph, text, **{k: v for k, v in spec.items() if k != "empty"})


def set_cell_text(cell, text: str, *, font_name: str, size_pt: float, bold: bool | None = False) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, font_name, size_pt, bold)


def format_table(table) -> None:
    if not table.rows:
        return
    header = table.rows[0].cells[0].text.strip()
    if header == "状況":
        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                font = FONT_GOTHIC if ri == 0 else FONT_MINCHO
                set_cell_text(cell, cell.text, font_name=font, size_pt=10, bold=False if ri else None)
    elif header == "順位":
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                size = 10 if ri == 0 or ci < 2 else 9
                set_cell_text(
                    cell,
                    cell.text,
                    font_name=FONT_GOTHIC,
                    size_pt=size,
                    bold=None if ri == 0 else False,
                )
    elif header.startswith("№"):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                size = 9 if ri > 0 and ci == 4 else 10
                set_cell_text(
                    cell,
                    cell.text,
                    font_name=FONT_GOTHIC,
                    size_pt=size,
                    bold=None if ri == 0 else False,
                )


def apply_document_formats(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            apply_paragraph_format(paragraph)
    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip():
            format_table(table)


def insert_paragraph_after(paragraph, text: str = "", *, center: bool = False) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if center:
        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return new_para


def insert_table_after(paragraph, doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    if doc.tables:
        table.style = doc.tables[0].style
    for ci, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[ci], header, font_name=FONT_GOTHIC, size_pt=10, bold=None)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            size = 9 if ci == 2 and len(headers) == 3 else (9 if ci == 4 else 10)
            set_cell_text(
                table.rows[ri].cells[ci],
                val,
                font_name=FONT_GOTHIC,
                size_pt=size,
                bold=False,
            )
    paragraph._p.addnext(table._tbl)
    return table
