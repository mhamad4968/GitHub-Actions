# -*- coding: utf-8 -*-
"""2026年度 情報セキュリティ勉強会テキスト（Word・本編）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "docs", "training", "security", "output", "情報セキュリティ勉強会テキスト_2026.docx")
OUT = os.path.abspath(OUT)

C_HEAD = RGBColor(0x1A, 0x36, 0x5D)
C_ACCENT = RGBColor(0xC0, 0x00, 0x00)
C_BLUE = RGBColor(0x2E, 0x5C, 0x8A)


def set_jp_font(run, name="游ゴシック", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_jp_font(run, size=sizes.get(level, 11), bold=True, color=C_HEAD if level <= 2 else C_BLUE)
    return p


def add_para(doc, text, bold=False, size=11, indent=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_jp_font(run, size=size, bold=bold)
    return p


def add_mixed_para(doc, parts, indent=0):
    """parts: list of (text, bold, color)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    for text, bold, color in parts:
        run = p.add_run(text)
        set_jp_font(run, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_jp_font(run)
    return p


def add_callout(doc, title, lines):
    """薄い背景の注意枠（表1行で代用）"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF5F5")
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    set_jp_font(r, bold=True, color=C_ACCENT, size=11)
    for line in lines:
        p2 = cell.add_paragraph()
        r2 = p2.add_run(line)
        set_jp_font(r2, size=10.5)
    doc.add_paragraph()


def add_numbered_steps(doc, steps):
    for i, step in enumerate(steps, 1):
        add_mixed_para(
            doc,
            [(f"{i}. ", True, C_ACCENT), (step, False, None)],
            indent=0.3,
        )


doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21.0)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

# ===== 表紙 =====
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("2026年度\n情報セキュリティ勉強会テキスト")
set_jp_font(r, size=22, bold=True, color=C_HEAD)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("システム推進室")
set_jp_font(r2, size=14)

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("実施期間：2026年7月15日〜12月15日（各支店訪問）\nテキスト研修 20分 ／ 動画視聴 20分")
set_jp_font(r3, size=11)

doc.add_page_break()

# ===== はじめに =====
add_heading(doc, "はじめに", 1)
add_para(
    doc,
    "本テキストは、毎年1回実施している情報セキュリティ勉強会用の資料です。"
    "社員の「情報セキュリティレベルの向上」および「当社の情報セキュリティに対する基本的な考え方の理解」を主な目的としています。",
)
add_para(doc, "【本日の流れ】", bold=True)
add_bullet(doc, "第1部　情報セキュリティの重要性（約5分）")
add_bullet(doc, "第2部　最近のセキュリティ動向（約15分）")
add_bullet(doc, "　　・偽セキュリティ警告（サポート詐欺）")
add_bullet(doc, "　　・ランサムウェア")
add_bullet(doc, "　　・AIを悪用した標的型メール")
add_bullet(doc, "第3部　動画視聴（20分）")
add_para(doc, "※ 理解度確認クイズは、別途 Microsoft Forms 等で実施予定です。", size=10)

# ===== 第1章 =====
add_heading(doc, "第1章　情報セキュリティの重要性", 1)

add_heading(doc, "1-1　情報セキュリティとは", 2)
add_para(doc, "「セキュリティ」は日本語では「安全」を意味します。英語には safety（セーフティ）と security（セキュリティ）の2語があり、意味が異なります。")
add_bullet(doc, "セキュリティ：侵入・盗難・破壊など、悪意をもって行われる人的な脅威に対する安全")
add_bullet(doc, "セーフティ：交通事故・自然災害など、偶発的・突発的な悪意のない脅威に対する安全")
add_mixed_para(
    doc,
    [
        ("以上より、", False, None),
        ("情報セキュリティ", True, C_BLUE),
        ("とは、", False, None),
        ("悪意をもって行われる第三者から、当社の情報（顧客情報・社員情報・業務情報等）を守ること", True, C_ACCENT),
        ("を指します。", False, None),
    ],
)

add_heading(doc, "1-2　なぜ今、全社員が学ぶ必要があるのか", 2)
add_para(
    doc,
    "サイバー攻撃のリスクは、企業の規模や業種を問わず国内全体に広がっています。"
    "警察庁の公表資料によれば、企業のサーバー等への攻撃に関する届出・相談は継続して発生しており、"
    "コンピュータウイルスの検知・感染、身代金を要求するサイバー攻撃（ランサムウェア）、"
    "脆弱性や設定不備を悪用した不正アクセス、ID・パスワードの窃取による不正アクセスなど、"
    "複数の類型が報告されています。",
)
add_para(
    doc,
    "2020年には、日経グループでウイルス付き電子メールにより1台のPCが感染し、"
    "約1万2,514人分の個人情報が流出した事例も報告されています。"
    "ウイルス検知システムを導入していても、新種の手口では検知が遅れる場合があります。",
    size=10.5,
)
add_callout(
    doc,
    "【重要】攻撃と対策は「いたちごっこ」です",
    [
        "ファイアウォールやウイルス対策ソフトなどのシステム対策だけでは不十分です。",
        "メールの1クリック、不審な画面への対応など、日常業務での判断が被害の有無を左右します。",
        "「社員一人ひとりの意識」が、会社全体の防衛線になります。",
    ],
)

# ===== 第2章 概観 =====
add_heading(doc, "第2章　最近のセキュリティ動向（IPA 10大脅威 2026）", 1)
add_para(
    doc,
    "独立行政法人情報処理推進機構（IPA）は、毎年「情報セキュリティ10大脅威」を公表しています。"
    "2026年版では、個人向け・組織向けそれぞれに、現在特に注意すべき脅威が整理されています。"
    "本年度の勉強会では、組織で特に問題となる次の3テーマを中心に学びます。",
)

# 表：10大脅威 組織向け 上位
add_heading(doc, "2-1　組織向け10大脅威（2026）上位と本研修の対応", 2)
table = doc.add_table(rows=5, cols=3)
table.style = "Table Grid"
hdr = ["順位", "脅威名称", "本テキストでの扱い"]
for i, h in enumerate(hdr):
    c = table.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        for run in p.runs:
            set_jp_font(run, bold=True, size=10)
rows = [
    ("1位", "ランサム攻撃による被害", "第2章 2-3 で詳述（11年連続選出）"),
    ("2位", "サプライチェーンや委託先を狙った攻撃", "委託先・取引先メールの確認を第2章 2-4 で言及"),
    ("3位", "AIの利用をめぐるサイバーリスク", "第2章 2-4 で詳述（2026年初選出）"),
    ("5位", "機密情報等を狙った標的型攻撃", "第2章 2-4 で詳述"),
]
for ri, row in enumerate(rows, 1):
    for ci, val in enumerate(row):
        table.rows[ri].cells[ci].text = val
        for p in table.rows[ri].cells[ci].paragraphs:
            for run in p.runs:
                set_jp_font(run, size=10)
doc.add_paragraph()

add_para(doc, "【個人向け】偽警告によるインターネット詐欺は、長年ランクインしており、IPA相談窓口にも「ウイルス検出の偽警告」に関する相談が多く寄せられています。業務PCでも発生しうるため、第2章 2-2 で重点的に学びます。")

# ===== 2-2 偽警告 =====
add_heading(doc, "2-2　偽セキュリティ警告（サポート詐欺）", 2)
add_heading(doc, "手口の概要", 3)
add_para(
    doc,
    "インターネット閲覧中に、ウイルス感染・システム破損・ライセンス期限切れなどを装った"
    "偽の警告画面（ポップアップ）が突然表示される手口です。"
    "画面には「今すぐ電話してください」「サポート番号」などが表示され、"
    "ユーザーに電話をかけさせます。",
)
add_para(doc, "電話口では、次のような要求が行われます。", bold=True)
add_bullet(doc, "ウイルス除去・サポート料金の支払い（電子マネー、ギフト券、プリペイドカード等）")
add_bullet(doc, "TeamViewer、AnyDesk 等の遠隔操作ソフトのインストール")
add_bullet(doc, "Windowsのイベントビューアやコマンドプロンプトでの操作")
add_bullet(doc, "「画面に映っている情報を見せてください」という指示（画面上の機密情報の窃取）")

add_heading(doc, "公表されている被害事例", 3)
add_para(
    doc,
    "2022年3月、国立研究開発法人国立がん研究センターは、職員がテレワーク中に"
    "偽警告画面の指示に従った結果、端末が第三者に乗っ取られ、"
    "ディスプレイに表示されていた被験者情報が閲覧された可能性があると公表しました。",
    size=10.5,
)
add_para(
    doc,
    "業務PCであっても、画面上に表示されている顧客情報・社員情報・医療・研究データ等は"
    "すべて漏えいリスクの対象となります。",
)

add_heading(doc, "被害にあわないための対応（必ず守ること）", 3)
add_callout(
    doc,
    "【絶対にやってはいけないこと】",
    [
        "警告画面に表示された電話番号に電話をかけること",
        "指示に従って遠隔操作ソフトをインストールすること",
        "指示に従ってコマンド入力や設定変更を行うこと",
        "電子マネー・ギフト券等で支払うこと",
    ],
)
add_numbered_steps(
    doc,
    [
        "警告画面が表示されたら、すぐにLANケーブルを抜く（Wi-Fi接続の場合はWi-Fiを切断する）。",
        "一人で対処しようとせず、上司・先輩社員に声をかける。",
        "システム推進室へ連絡する（連絡手段は社内掲示・マニュアルに従う）。",
        "指示を受けるまで、該当PCで業務を再開しない。",
    ],
)
add_para(doc, "※ 画面の閉じ方は、動画「偽セキュリティ警告画面に注意！【画面削除方法】」でも確認します。", size=10)

add_heading(doc, "講師メモ（話すポイント）", 3)
add_para(
    doc,
    "・「Microsoftやウイルス対策ソフトの正規サポートは、画面に電話番号を表示して電話をかけさせることはない」"
    "と強調してください。"
    "・業務中に個人のWeb閲覧（検索・動画・SNS等）でも発生しうることを伝えてください。"
    "・「警告画面を閉じる操作」は動画で実演するため、テキストでは初動（切断・報告）を優先して説明します。",
    size=10.5,
)

# ===== 2-3 ランサムウェア =====
add_heading(doc, "2-3　ランサムウェア（身代金型マルウェア）", 2)
add_heading(doc, "ランサムウェアとは", 3)
add_para(
    doc,
    "ランサムウェア（Ransomware）は、端末内のファイルを暗号化し、復号と引き換えに金銭（身代金）を要求する不正プログラムです。"
    "近年は、暗号化に加えてデータを窃取し、公開すると脅す「二重脅迫」が一般化しています。"
    "さらに、DDoS攻撃の実施や、被害者の顧客・取引先への連絡など、"
    "「三重脅迫」「四重脅迫」と呼ばれる手口も確認されています。",
)
add_para(
    doc,
    "DDoS攻撃（Distributed Denial of Service）とは、複数の機器を踏み台にして"
    "特定のサーバーへ大量のアクセスを集中させ、サービスを停止させる攻撃です。",
    size=10.5,
)

add_heading(doc, "主な感染経路", 3)
add_para(doc, "警察庁・IPA等の公表資料では、主に次の経路が報告されています。")
add_bullet(doc, "標的型攻撃メール・迷惑メールの添付ファイル、または本文中のURLをクリック")
add_bullet(doc, "VPN機器・リモートデスクトップ等の脆弱性・設定不備・漏えいした認証情報の悪用")
add_bullet(doc, "不正なUSBメモリ等の記録媒体")
add_bullet(doc, "RaaS（Ransomware as a Service）：攻撃ツールの提供・攻撃代行サービスの悪用")
add_para(doc, "※ 感染経路のグラフは、警察庁「サイバー空間をめぐる脅威の情勢等」資料を参照してください。", size=10)

add_heading(doc, "事例：KADOKAWAグループ（2024年6月）", 3)
add_para(
    doc,
    "2024年6月8日、KADOKAWAグループおよびニコニコ動画等で大規模なランサムウェア攻撃が発生しました。"
    "出版事業の出荷部数が平常の3分の1に減少、動画サービスの停止、オンラインショップへのログイン不可など、"
    "社会に広く影響する事態となりました。",
)
add_para(doc, "報道・公表情報によると、次の点が教訓として挙げられています。", bold=True)
add_bullet(doc, "約1.5TB規模のデータ窃取が報じられた（個人情報、業務マニュアル、契約書等）")
add_bullet(doc, "特定の担当者が業務上必要以上に個人フォルダへデータをコピーしたまま放置し、漏えい範囲が拡大した")
add_bullet(doc, "復旧・調査に長期間を要し、事業継続に重大な影響が出た")
add_callout(
    doc,
    "【当社で守ること】",
    [
        "重要情報を個人フォルダにコピーして放置しない（必要最小限・使用後は削除）。",
        "共有フォルダ・業務システム上の正規の保管場所を使う。",
        "身代金を支払っても復号・再攻撃が保証されないことを理解する。",
    ],
)

add_heading(doc, "被害が起きた組織に起こること", 3)
add_bullet(doc, "業務システム・ファイルサーバーが使用不能になり、出荷・請求・顧客対応が止まる")
add_bullet(doc, "復旧作業に数日〜数週間かかり、正常化までの間、代替手段での業務が必要になる")
add_bullet(doc, "個人情報・取引先情報の漏えいが公表されると、信用失墜・損害賠償・行政指導のリスクがある")
add_bullet(doc, "身代金を支払っても、データが復号されない・再度攻撃される事例も報告されている")
add_para(
    doc,
    "2025年には、国内でも大規模企業を対象としたランサムウェア攻撃により、"
    "グループ会社を経由した被害拡大や、電子取引サービス（EOS）停止など、"
    "サプライチェーン全体に波及した事例も公表されています（IPA 10大脅威2026 解説書参照）。",
    size=10.5,
)

# ===== 2-4 AI =====
add_heading(doc, "2-4　AIを悪用した標的型メール", 2)
add_para(
    doc,
    "IPA「情報セキュリティ10大脅威 2026」では、「AIの利用をめぐるサイバーリスク」が組織向け3位に初選出されました。"
    "生成AI（ChatGPT等）の普及により、従来より文法の正しい、自然な日本語のフィッシングメールが"
    "短時間・低コストで作成できるようになっています。",
)
add_heading(doc, "従来のフィッシングとの違い", 3)
add_bullet(doc, "不自然な日本語・変な敬語が少なく、一見すると本物のビジネスメールに見える")
add_bullet(doc, "上司名・取引先名・社内用語を含めたなりすましが容易")
add_bullet(doc, "大量のバリエーションを自動生成し、標的型攻撃の精度が上がる")
add_bullet(doc, "音声・画像によるなりすまし（ディープフェイク）と組み合わせた手口も報告されている")

add_heading(doc, "典型的な手口", 3)
add_bullet(doc, "「至急」「稟議なし」「社外秘」等、急かす件名・本文")
add_bullet(doc, "「振込先が変更になりました」等のビジネスメール詐欺（BEC）")
add_bullet(doc, "「パスワードの確認」「セキュリティ更新のためクリック」等を装うリンク")
add_bullet(doc, "短縮URLや似せたドメイン名（例：j-bis.co.jp に似た文字列）")

add_heading(doc, "見分け方・対策", 3)
add_numbered_steps(
    doc,
    [
        "送信者アドレス・件名・本文の言い回しが、いつもと違う場合は開封・リンククリック・添付ファイルの開封をしない。",
        "「至急」「個別対応」で振込・認証情報の入力を求められた場合、メールだけで判断せず、電話等で送信者本人を確認する。",
        "URLはメール内リンクをクリックせず、ブラウザに直接入力するか、ブックマーク・公式サイトからアクセスする。",
        "不審なメールは削除せず、システム推進室へ報告する（類似メールの社内周知に役立つ）。",
    ],
)

add_heading(doc, "具体例：こんなメールに注意", 3)
examples = [
    "件名「【至急】請求書の送付先変更の件」— いつもの取引先名だが、メールアドレスのドメインが1文字違う。",
    "件名「重要：パスワード期限切れ」— 社内のIT部門を装っているが、リンク先が j-bis-co-jp.com 等の偽サイト。",
    "本文「稟議不要ですので本日中に処理を」— 上司の名前が入っているが、普段使わない言い回し。",
    "添付ファイル「請求書_202607.pdf.exe」— 拡張子を二重にして見せかけている。",
]
for ex in examples:
    add_bullet(doc, ex)

add_heading(doc, "ビジネスメール詐欺（BEC）について", 3)
add_para(
    doc,
    "IPA組織向け10大脅威にも「ビジネスメール詐欺」がランクインしています。"
    "取引先や上司になりすまし、「振込先口座が変わった」「急ぎで振込してほしい」"
    "とメールで指示し、本来の口座とは別の口座へ送金させる手口です。"
    "AIにより、より自然な文面のメールが作成されるようになっています。",
)
add_callout(
    doc,
    "【振込・支払いに関するメールを受け取ったら】",
    [
        "メールだけで判断せず、電話等で相手先担当者に直接確認する（メールに書いてある連絡先は使わない）。",
        "口座変更の指示は、社内規定の承認フロー・経理担当を必ず通す。",
        "「至急」「秘密」で個人判断を迫るメールは、より慎重に確認する。",
    ],
)

add_heading(doc, "講師メモ（話すポイント）", 3)
add_para(
    doc,
    "・「AIだから見分けられない」ではなく、「いつもと違う」「急かす」「金銭・認証情報」を"
    "キーワードに止まる、と伝えてください。"
    "・社内メールでも、なりすまし・乗っ取りは起こりうることを説明してください。",
    size=10.5,
)

# ===== 第3章 当社ルール =====
add_heading(doc, "第3章　当社で守るべき基本ルール", 1)
add_para(doc, "以下は、これまでの勉強会でも共有している当社の基本ルールです。日常業務で必ず守ってください。")

rules = [
    (
        "メール・Web",
        "見覚えのないメールは開封しない。添付ファイル・リンクは絶対に開かない。",
    ),
    (
        "記録媒体",
        "USBメモリ等は会社で購入・管理したもののみ使用する。私物の利用は禁止。",
    ),
    (
        "ネットワーク",
        "自宅・出張先から業務する場合、会社支給のポケットWi-Fi等以外に接続しない。公共のフリーWi-Fiは利用しない。",
    ),
    (
        "異常時の初動",
        "端末の動作がおかしい、警告画面が出た、不審なメールを開いてしまった等 → ①LAN/Wi-Fiを切断 ②システム推進室へ連絡。",
    ),
    (
        "情報の保管",
        "顧客・社員等の重要情報を、個人フォルダに不必要にコピーして放置しない。",
    ),
]
for title, body in rules:
    add_mixed_para(doc, [(f"■ {title}：", True, C_BLUE), (body, False, None)])

add_heading(doc, "3-1　パスワード・アカウント管理", 2)
add_bullet(doc, "社内規定に従い、推測されにくいパスワードを設定し、使い回しをしない")
add_bullet(doc, "多要素認証（MFA）が利用できるサービスでは、可能な限り有効にする")
add_bullet(doc, "パスワードやワンタイムコードを、メール・チャットで他人に送らない")
add_bullet(doc, "退職者・異動者のアカウントは、人事・システム推進室の手続きに従い速やかに無効化する")

add_heading(doc, "3-2　よくある誤解（Q&A）", 2)
faqs = [
    (
        "Q. ウイルス対策ソフトが入っているから大丈夫では？",
        "A. 最新の手口（ゼロデイ、フィッシング、偽警告）は検知できない場合があります。ソフトは「補助線」であり、"
        "メール・Web・USBの取り扱いは自分で判断する必要があります。",
    ),
    (
        "Q. 個人のスマホや自宅PCなら会社に関係ない？",
        "A. 業務メール・クラウドサービスへのアクセス、VPN接続等を行う場合、"
        "会社の情報を扱うことになります。私物端末の業務利用は社内規定に従ってください。",
    ),
    (
        "Q. 警告画面が出たが、閉じられなかった。",
        "A. 電話はかけず、LAN/Wi-Fiを切断し、システム推進室へ連絡してください。"
        "電源ボタン長押しで強制終了する場合も、再起動後は業務再開前に必ず報告を。",
    ),
    (
        "Q. 不審なメールを開いてしまった。",
        "A. 添付ファイルを開いていなければ被害に至らない場合もありますが、"
        "リンクをクリックした・情報を入力した場合は、ただちに切断・報告してください。",
    ),
]
for q, a in faqs:
    add_mixed_para(doc, [(q + "\n", True, C_ACCENT), (a, False, None)])

# ===== 第4章 動画 =====
add_heading(doc, "第4章　動画視聴（20分）", 1)
add_para(doc, "テキスト研修の後、以下の3本を視聴してください。内容は本テキスト第2章と対応しています。")

videos = [
    ("① そこにある脅威〜組織を狙うランサムウェア攻撃〜", "https://youtu.be/TWqJ5P8oaUM?t=8", "ランサムウェアの脅威（2-3）"),
    ("② 偽セキュリティ警告画面（サポート詐欺）に注意！【画面削除方法】", "https://youtu.be/BV6GebTEiQI", "偽警告の対処（2-2）"),
    ("③ 見えざるサイバー攻撃", "https://www.youtube.com/watch?v=ZZJ7VMJ5Btw", "サイバー攻撃全般の理解"),
]
for title, url, note in videos:
    add_para(doc, title, bold=True)
    add_para(doc, url, size=10)
    add_para(doc, f"（対応：{note}）", size=10)
    doc.add_paragraph()

add_callout(
    doc,
    "【視聴後】",
    [
        "不明点があれば、システム推進室へ質問してください。",
        "クイズは別途 Microsoft Forms 等で実施します。",
    ],
)

# ===== おわり =====
add_heading(doc, "おわりに", 1)
add_para(
    doc,
    "サイバー攻撃の手口は日々変化していますが、「怪しいと思ったら自分で処理しない」"
    "「ネットワークを切り離す」「すぐに報告する」という基本動作が、"
    "自分・顧客・会社を守る最後の防衛線になります。"
    "本テキストで学んだ内容を、日々の業務の中で実践してください。",
)
add_para(doc, "以上", bold=True)
t_end = doc.add_paragraph()
t_end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_end = t_end.add_run("システム推進室")
set_jp_font(r_end, size=11)

# ===== 参考資料 =====
doc.add_page_break()
add_heading(doc, "参考資料・出典", 1)
refs = [
    "独立行政法人情報処理推進機構（IPA）「情報セキュリティ10大脅威 2026」",
    "警察庁「サイバー空間をめぐる脅威の情勢等について」",
    "日本経済新聞社 プレスリリース「サイバー攻撃による社員等の個人情報流出について」（2020年）",
    "国立研究開発法人国立がん研究センター 公表資料（2022年）",
    "KADOKAWAグループ インシデント関連公表（2024年）",
]
for ref in refs:
    add_bullet(doc, ref)

add_para(doc, "【差し込み用イメージ（PowerPoint化する場合）】", bold=True, size=10)
add_bullet(doc, "偽警告画面のスクリーンショット")
add_bullet(doc, "ランサムウェア感染画面のイメージ")
add_bullet(doc, "警察庁資料：ランサムウェア感染経路グラフ")

add_heading(doc, "付録A　組織向け10大脅威2026（一覧）", 1)
table2 = doc.add_table(rows=11, cols=2)
table2.style = "Table Grid"
table2.rows[0].cells[0].text = "順位"
table2.rows[0].cells[1].text = "脅威名称"
threats_full = [
    "1", "ランサム攻撃による被害",
    "2", "サプライチェーンや委託先を狙った攻撃",
    "3", "AIの利用をめぐるサイバーリスク",
    "4", "システムの脆弱性を悪用した攻撃",
    "5", "機密情報等を狙った標的型攻撃",
    "6", "地政学的リスクに起因するサイバー攻撃（情報戦を含む）",
    "7", "内部不正による情報漏えい等",
    "8", "リモートワーク等の環境や仕組みを狙った攻撃",
    "9", "DDoS攻撃（分散型サービス妨害攻撃）",
    "10", "ビジネスメール詐欺",
]
for i in range(10):
    table2.rows[i + 1].cells[0].text = threats_full[i * 2]
    table2.rows[i + 1].cells[1].text = threats_full[i * 2 + 1]
add_para(doc, "出典：IPA「情報セキュリティ10大脅威 2026」組織編", size=9)

add_heading(doc, "付録B　20分テキスト研修 進行目安", 1)
schedule = [
    ("0:00〜5:00", "第1章　情報セキュリティの重要性"),
    ("5:00〜8:00", "第2章 2-1　IPA 10大脅威の概観"),
    ("8:00〜11:00", "第2章 2-2　偽セキュリティ警告"),
    ("11:00〜14:00", "第2章 2-3　ランサムウェア"),
    ("14:00〜18:00", "第2章 2-4　AI標的型メール"),
    ("18:00〜20:00", "第3章　当社ルール・Q&A"),
]
t3 = doc.add_table(rows=len(schedule) + 1, cols=2)
t3.style = "Table Grid"
t3.rows[0].cells[0].text = "時間"
t3.rows[0].cells[1].text = "内容"
for i, (tm, content) in enumerate(schedule, 1):
    t3.rows[i].cells[0].text = tm
    t3.rows[i].cells[1].text = content

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("Saved:", OUT)
# 文字数概算
from docx import Document as D2
d = D2(OUT)
chars = sum(len(p.text) for p in d.paragraphs)
for table in d.tables:
    for row in table.rows:
        for cell in row.cells:
            chars += sum(len(p.text) for p in cell.paragraphs)
print("Approx chars:", chars)
