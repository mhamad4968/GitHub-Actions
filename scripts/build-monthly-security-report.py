# -*- coding: utf-8 -*-
"""Build monthly 情報セキュリティレポート DOCX (R1)."""
from __future__ import annotations

import argparse
import json
import shutil
import sys
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from matplotlib import font_manager

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from lib import docx_template_format as dtf  # noqa: E402

IPA_HEADER = ("順位", "脅威", "特に注意すべきポイント")

CHART_SPECS = [
    {
        "filename": "chart-victim-scale.png",
        "title": "被害報告の規模別内訳",
        "labels": ["中小企業\n63%", "大企業\n28%", "団体等\n9%"],
        "values": [63, 28, 9],
        "colors": ["#4472C4", "#ED7D31", "#A5A5A5"],
    },
    {
        "filename": "chart-business-impact.png",
        "title": "業務への影響",
        "labels": ["一部以上の\n業務に影響\n81%", "全業務停止\n10%", "影響なし\n9%"],
        "values": [81, 10, 9],
        "colors": ["#C00000", "#FFC000", "#70AD47"],
    },
    {
        "filename": "chart-recovery-period.png",
        "title": "復旧に要した期間",
        "labels": ["1週間未満\n26%", "1週間～1ヶ月\n25%", "1～2ヶ月\n17%", "2ヶ月以上\n7%", "復旧中\n25%"],
        "values": [26, 25, 17, 7, 25],
        "colors": ["#5B9BD5", "#4472C4", "#264478", "#1F3864", "#BDD7EE"],
    },
    {
        "filename": "chart-backup-restore.png",
        "title": "バックアップからの復元結果\n（取得済み企業）",
        "labels": ["直前水準まで\n復元不能\n74%", "復元できた\n26%"],
        "values": [74, 26],
        "colors": ["#C00000", "#70AD47"],
    },
]


def setup_japanese_font() -> str:
    for name in ("MS Gothic", "ＭＳ ゴシック", "Yu Gothic", "Meiryo"):
        if name in {f.name for f in font_manager.fontManager.ttflist}:
            plt.rcParams["font.family"] = name
            plt.rcParams["axes.unicode_minus"] = False
            return name
    plt.rcParams["axes.unicode_minus"] = False
    return "default"


def make_pie_chart(spec: dict, out_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(4.2, 3.4), dpi=150)
    ax.pie(
        spec["values"],
        labels=spec["labels"],
        colors=spec["colors"],
        startangle=90,
        counterclock=False,
        autopct="",
        labeldistance=1.08,
        wedgeprops={"linewidth": 0.6, "edgecolor": "white"},
    )
    for label in ax.texts:
        label.set_fontsize(8)
    ax.set_title(spec["title"], fontsize=10, fontweight="bold", pad=10)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_cost_bar_chart(out_path: Path) -> None:
    labels = ["100万円\n未満", "100～500\n万円", "500～1000\n万円", "1000～5000\n万円", "5000万～\n1億円", "1億円\n以上"]
    values = [23, 16, 12, 27, 15, 8]
    colors = ["#5B9BD5", "#4472C4", "#264478", "#C00000", "#ED7D31", "#843C0C"]
    fig, ax = plt.subplots(figsize=(5.2, 3.2), dpi=150)
    bars = ax.bar(labels, values, color=colors, edgecolor="white", linewidth=0.6)
    ax.set_title("調査・復旧にかかった費用（内訳）", fontsize=10, fontweight="bold", pad=10)
    ax.set_ylabel("割合（％）", fontsize=9)
    ax.set_ylim(0, 32)
    ax.tick_params(axis="both", labelsize=8)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.6, f"{val}%", ha="center", va="bottom", fontsize=8)
    fig.tight_layout()
    fig.savefig(out_path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def generate_charts(charts_dir: Path) -> list[Path]:
    charts_dir.mkdir(parents=True, exist_ok=True)
    setup_japanese_font()
    paths = []
    for spec in CHART_SPECS:
        path = charts_dir / spec["filename"]
        make_pie_chart(spec, path)
        paths.append(path)
    cost = charts_dir / "chart-recovery-cost.png"
    make_cost_bar_chart(cost)
    paths.append(cost)
    return paths


def insert_chart_grid_after(paragraph, doc, chart_paths: list[Path], cols: int = 2) -> None:
    rows = (len(chart_paths) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    if doc.tables:
        table.style = doc.tables[0].style
    for idx, path in enumerate(chart_paths):
        cell = table.rows[idx // cols].cells[idx % cols]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(2.55))
    paragraph._p.addnext(table._tbl)


def load_config(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def build_report(cfg: dict) -> Path:
    work_dir = Path(cfg["work_dir"])
    work_dir.mkdir(parents=True, exist_ok=True)
    template = next(work_dir.glob(cfg["template_glob"]))
    out = work_dir / cfg["output_filename"]
    charts_dir = work_dir / "_charts"

    if not cfg.get("detection_confirmed", True):
        print("[preflight] detection_confirmed=false — §2 数値はプレースホルダの可能性あり（R3）")

    shutil.copy2(template, out)
    doc = Document(out)
    paras = doc.paragraphs
    sec2_idx = next(i for i, p in enumerate(paras) if p.text.startswith("２.２０２６年"))

    title = f"２０２６年{cfg['report_month_label']}月情報セキュリティレポート"
    dtf.set_paragraph_text(paras[0], title, size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    dtf.set_paragraph_text(paras[1], cfg["meeting_date"], align=WD_ALIGN_PARAGRAPH.RIGHT)
    dtf.set_paragraph_text(paras[2], "経営企画部システム推進室", align=WD_ALIGN_PARAGRAPH.RIGHT)

    start = 3
    for i, text in enumerate(cfg["section1"]):
        if start + i < sec2_idx:
            dtf.set_paragraph_text(paras[start + i], text)
    for i in range(start + len(cfg["section1"]), sec2_idx):
        dtf.clear_paragraph(paras[i])

    graph_heading = next(
        (p for p in doc.paragraphs if p.text.strip() == "●ランサムウェア被害の実態（参考グラフ・警視庁令和6年統計より）"),
        None,
    )
    if graph_heading is None:
        raise SystemExit("Graph heading not found")
    insert_chart_grid_after(graph_heading, doc, generate_charts(charts_dir), cols=2)
    graph_tbl = graph_heading._p.getnext()
    while graph_tbl is not None and not graph_tbl.tag.endswith("tbl"):
        graph_tbl = graph_tbl.getnext()
    if graph_tbl is not None:
        cap_p = OxmlElement("w:p")
        graph_tbl.addnext(cap_p)
        cap_para = Paragraph(cap_p, graph_heading._parent)
        dtf.set_paragraph_text(
            cap_para,
            "（出典：警視庁「令和6年サイバー空間をめぐる脅威の情勢等について」／0024792412.pdf を基にシステム推進室が作成）",
        )

    ipa_heading = next(
        (p for p in doc.paragraphs if p.text.strip() == "●IPA 2026年版で特に注意すべきポイント"),
        None,
    )
    if ipa_heading is None:
        raise SystemExit("IPA heading not found")
    dtf.insert_table_after(ipa_heading, doc, IPA_HEADER, [tuple(r) for r in cfg["ipa_threat_rows"]])

    for i, text in enumerate(cfg["section2"]):
        if sec2_idx + i < len(paras):
            dtf.set_paragraph_text(paras[sec2_idx + i], text)
    for i in range(sec2_idx + len(cfg["section2"]), len(paras)):
        dtf.clear_paragraph(paras[i])

    for table in doc.tables:
        if table.rows and table.rows[0].cells[0].text.strip().startswith("№"):
            for ri, case in enumerate(cfg["external_cases"], start=1):
                if ri < len(table.rows):
                    for ci, val in enumerate(case):
                        size = 9 if ci == 4 else 10
                        dtf.set_cell_text(
                            table.rows[ri].cells[ci],
                            val,
                            font_name=dtf.FONT_GOTHIC,
                            size_pt=size,
                            bold=False,
                        )
            break

    dtf.apply_document_formats(doc)
    doc.save(out)
    return out


def main() -> None:
    parser = argparse.ArgumentParser(description="Build monthly security report DOCX")
    parser.add_argument(
        "--config",
        type=Path,
        default=SCRIPT_DIR / "data" / "monthly-security-report-202605.json",
        help="JSON config path",
    )
    args = parser.parse_args()
    out = build_report(load_config(args.config))
    print("OK", out)


if __name__ == "__main__":
    main()
